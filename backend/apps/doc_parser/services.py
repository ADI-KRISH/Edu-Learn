import os
from google.cloud import documentai_v1 
from django.conf import settings
from urllib3 import response
from  .models import Document

class DocumentParser:
    def __init__(self):
        self.client = documentai_v1.DocumentProcessorServiceClient()
        self.project_id = settings.GOOGLE_PROJECT_ID
        self.processor_id = settings.DOCAI_PROCESSOR_ID
    def determine_f_type(self,filename:str) -> str:
        ext = os.path.splitext(filename)[1].lower()
        if ext == '.pdf':
            return pdf
        elif ext in ['.jpg','.jpeg','.png']:
            return 'image'
        elif ext == '.docx' :
            return 'docx'
        raise ValueError('Unsupported file type')
    def process_doc(self,document:Document) -> str:
        try:
            dp = document.original_file.path
            f_t = self,determine_f_t(doc_path)
            if  f_t == 'pdf':
                return self.process_pdf(dp)
            elif  f_t == 'image':
                return self.process_image(dp)
            elif f_t == 'docx' :
                return self.process_docx(dp)
        except Exception as e:
            document.status = 'failed'
            document.save()
            raise e
    def process_pdf(self, fp:str) ->str:
        with open(file_path,'rb') as f:
            content = f.read()
            request = documentai_v1.ProcessRequest(
                name = f'projects/{self.project_id}/locations/us/processors/{self.processor_id}',
                raw_doc = documentai_v1.RawDocument(
                    content = content,
                    mime_type = 'application/pdf'
                )  
            )
            res = self.client.process_document(requenst = request)
            return self.document.text
    def process_image(self,fp:str) ->str:
        from google.cloud import vision
        client = vision.ImageAnnotatorClient()
        with open(fp,'rb') as f:
            content = f.read()
            image = vision.Image(content=content)
            return response.text_annotations[0].description
    def process_docx(self,fp:str) ->str:
        from docx import Document
        doc = Document(fp)
        return '\n'.join([para.text for para in doc.paragraphs])